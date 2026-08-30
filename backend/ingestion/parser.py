from pathlib import Path

import pymupdf
from docx import Document
from openpyxl import load_workbook

from backend.memory.memory_item import MemoryItem


def extract_pdf(path):

    doc = pymupdf.open(path)

    text = ""

    for page in doc:
        text += page.get_text() + "\n"

    metadata = {
        "pages": len(doc)
    }

    return text, metadata


def extract_docx(path):

    doc = Document(path)

    text = ""

    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    metadata = {
        "paragraphs": len(doc.paragraphs)
    }

    return text, metadata


def extract_txt(path):

    # Most files are UTF-8, but a raw UnicodeDecodeError on
    # anything else (Windows-exported notes, older files, etc.)
    # used to crash the whole ingest with no fallback. Try the
    # common cases in order; latin-1 always succeeds last since
    # it maps every possible byte value, so this never raises.

    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):

        try:

            text = Path(path).read_text(
                encoding=encoding
            )

            return text, {"encoding": encoding}

        except UnicodeDecodeError:

            continue

    # Unreachable in practice (latin-1 always succeeds), kept
    # only as an explicit fallback rather than falling through
    # silently.

    text = Path(path).read_text(
        encoding="latin-1"
    )

    return text, {"encoding": "latin-1"}


def extract_xlsx(path):

    workbook = load_workbook(
        path,
        read_only=True,
        data_only=True
    )

    text = ""

    for sheet in workbook.worksheets:

        text += f"\n--- SHEET: {sheet.title} ---\n"

        for row in sheet.iter_rows(values_only=True):

            values = [
                str(value)
                for value in row
                if value is not None
            ]

            if values:
                text += " | ".join(values) + "\n"

    metadata = {
        "sheets": workbook.sheetnames
    }

    return text, metadata


def parse_file(path):

    path = Path(path)

    extension = path.suffix.lower()

    if extension == ".pdf":

        text, metadata = extract_pdf(path)
        modality = "document"

    elif extension == ".docx":

        text, metadata = extract_docx(path)
        modality = "document"

    elif extension in [".txt", ".md"]:

        text, metadata = extract_txt(path)
        modality = "text"

    elif extension == ".xlsx":

        text, metadata = extract_xlsx(path)
        modality = "spreadsheet"

    else:

        raise ValueError(
            f"Unsupported file type: {extension}"
        )

    return MemoryItem(
        id=path.stem,
        source=str(path),
        modality=modality,
        content=text,
        metadata=metadata
    )


if __name__ == "__main__":

    file_path = input("Enter file path: ")

    memory = parse_file(file_path)

    print("\n========== MEMORY ITEM ==========\n")

    print("ID:", memory.id)
    print("Source:", memory.source)
    print("Modality:", memory.modality)
    print("Metadata:", memory.metadata)

    print("\n---------- CONTENT ----------\n")

    print(memory.content[:5000])

    print("\n================================\n")